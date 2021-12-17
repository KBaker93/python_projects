from docx import Document
from docx.shared import Inches

document = Document()


# picture: Ensure that picture being used is included in folder and that name of the photo is used in the string.

document.add_picture(
    'me.jpg', 
    width = Inches(1.0))
# Contact information

name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')
social_media = input('What is your social media link? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email + ' | ' + social_media + ' | ')

# about me

document.add_heading('About Me')
about_me = input('Tell me about yourself ')
document.add_paragraph(about_me)

# Work experience. 
document.add_heading('Work Experience ')
p = document.add_paragraph()

role = input('Enter Role tite ')
from_date = input('From date ')
to_date = input('to date ')
company = input('Company ')

p.add_run(role + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True
p.add_run(company + '\n') 

experience_details = input(
    'Describe your experience at ' + company )
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? yes or no ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        role = input('Enter Role tite ')
        from_date = input('From date ')
        to_date = input('to date ')
        company = input('Company ')

        p.add_run(role + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True
        p.add_run(company + '\n') 

        experience_details = input(
         'Describe your experience at ' + company)
        p.add_run(experience_details)

    else:
        break

# skills 
document.add_heading('Skills')
skill = input('enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input(
        'Do you have more skills? yes or no ')
    if has_more_skills.lower() == 'yes':
        skill = input('enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# projects
document.add_heading('Achievements')
p = document.add_paragraph()

achievement = input('achievement title ')
p.add_run(achievement + '\n').bold = True

achievement_details = input('Describe your achievement ')
p.add_run(achievement_details + '\n')

# more achievements. Make sure to add break to end loop
while True:
    has_more_achievements = input(
        'Do you have more achievements? yes or no ')
    if has_more_achievements.lower() == 'yes':
        p = document.add_paragraph()

        achievement = input('achievement title')
        p.add_run(achievement + '\n').bold = True

        achievement_details = input('Describe your achievement')
        p.add_run(achievement_details + ' ')
    else:
        break



document.save('cv.docx')